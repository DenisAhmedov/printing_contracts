from docx import Document
from docx.shared import Pt
from python_docx_replace import docx_replace
import os
import sys
import shutil
from datetime import datetime
from tqdm import tqdm
import time

file = os.path.join(os.getcwd(), 'template.docx')
temp_dir_path = os.path.join(os.getcwd(), 'temp')
output_file = os.path.join(temp_dir_path, 'output.docx')
year = datetime.strftime(datetime.now(), '%y')
full_year = datetime.strftime(datetime.now(), '%Y')
docx_files = []

if not os.path.exists(file):
    sys.exit('Не обнаружен файл template.docx')

if not os.path.exists(temp_dir_path):
    os.mkdir(temp_dir_path)

# Удаление файлов из временной директории
for files in os.listdir(temp_dir_path):
    temp_file = os.path.join(temp_dir_path, files)
    try:
        shutil.rmtree(temp_file)
    except OSError:
        os.remove(temp_file)

if (len(sys.argv) == 4) and (int(sys.argv[2]) >= int(sys.argv[1])):
    start_num = int(sys.argv[1])
    end_num = int(sys.argv[2])
    prefix = sys.argv[3]
else:
    start_num = int(input('Введите начальный номер договора: '))
    end_num = int(input('Введите конечный номер договора: '))
    prefix = input('Введите префикс: ')
    if start_num > end_num:
        exit('Начальный номер не может быть больше конечного')


# Создание временных файлов
for num in tqdm(range(start_num, end_num + 1)):
    doc = Document(file)
    contract_number_string = f'{prefix}-{num:03}-{year}'
    docx_replace(doc, contract_number = contract_number_string, year = full_year)

    for section in doc.sections:
        section.header.is_linked_to_previous = True


    temp_file = os.path.join(temp_dir_path, f'{num:03}.docx')
    doc.save(temp_file)
    time.sleep(0.1)

for filedocx in os.listdir(temp_dir_path):
    if filedocx.endswith('.docx'):
        docx_files.append(os.path.join(temp_dir_path, filedocx))

merged_document = Document(docx_files[0])

# Склеивание временных файлов в один итоговый
for index, file in enumerate(docx_files):
    sub_doc = Document(file)
    for element in sub_doc.element.body:
        merged_document.element.body.append(element)

    if index != 0:
        sub_doc = Document(file)
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)


merged_document.save(output_file)
time.sleep(0.5)
os.startfile(output_file, 'print')
